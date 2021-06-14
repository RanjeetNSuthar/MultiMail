import os
import re
from docx import Document
import openpyxl
from docx2pdf import convert

attachment_list = []

class replacement:
    def __init__(self,sh1=None,doc_obj=None):
        self.sh1 = sh1
        self.doc_obj = doc_obj


    @staticmethod
    def docx_replace_regex(doc_obj, regex, replace):
        rep_obj = replacement()
        print("inside replacement")

        for p in doc_obj.paragraphs:

            if regex.search(p.text):
                inline = p.runs
                # Loop added to work with runs (strings with same style)
                for i in range(len(inline)):
                    if '{' in inline[i - 1].text:
                        regex_without_braces = re.compile(str(regex.pattern)[1:-1])
                        # print(inline[i].text)
                        if regex_without_braces.search(inline[i].text):
                            text = regex_without_braces.sub(replace, inline[i].text)  # replace the leftmost occurrence of {regex_without_braces} pattern by {replace} in string {inline[i].text}
                            inline[i].text = text
                            inline[i].italic = False
                            inline[i - 1].text = str(inline[i - 1].text).replace('{', '')
                            inline[i + 1].text = str(inline[i + 1].text).replace('}', '')

        for table in doc_obj.tables:
            for row in table.rows:
                for cell in row.cells:
                    rep_obj.docx_replace_regex(cell, regex, replace)




    def creating_dictionary(self):
        dictionary = {}

        for columns in range(1, self.sh1.max_column + 1):
            column_head_value = '{' + self.sh1.cell(1, columns).value + '}'
            dictionary.update({column_head_value: None})

        key_list = list(dictionary.keys())

        for rows in range(2, self.sh1.max_row + 1):

            for column in range(1, self.sh1.max_column + 1):
                dictionary[str(key_list[column - 1])] = str(self.sh1.cell(rows, column).value)

            for word, replacement in dictionary.items():
                word_re = re.compile(word)
                self.docx_replace_regex(self.doc_obj, word_re, replacement)

            filename2 = str(self.sh1.cell(rows, 1).value) + str(self.sh1.cell(rows, 2).value) +".pdf"
            filename3 = str(self.sh1.cell(rows, 1).value) + str(self.sh1.cell(rows, 2).value)+".docx"
            attachment_list.append(filename2)
            self.doc_obj.save("../miniproject/Doc/"+filename3)
            convert("../miniproject/Doc/"+filename3)
            os.remove("../miniproject/Doc/"+filename3)
            self.doc_obj = Document('../miniproject/Template.docx')


    @staticmethod
    def updating_attachment_list(attachment_list):
        wb = openpyxl.load_workbook("../miniproject/email&attachment_list.xlsx")
        sh6 = wb['Sheet1']
        sh6.delete_rows(2, sh6.max_row)
        for rows in range(2, (len(attachment_list) + 2)):
            attachment_no = rows - 2
            sh6.cell(rows, 4).value = "../miniproject/Doc/" + str(attachment_list[attachment_no])
            wb.save('../miniproject/email&attachment_list.xlsx')

class Creating_record_file:
    def __init__(self,doc_obj,sheet_obj):
        self.doc_obj = doc_obj
        self.sheet1 = sheet_obj

    def Creating_record_headings(self):
        print("inside heading")

        list_of_record_headings = []
        self.sheet1.delete_cols(1, self.sheet1.max_column)
        self.sheet1.delete_rows(1, self.sheet1.max_row)

        regex = re.compile(r"([{])([a-zA-Z]*)([}])")

        for p in self.doc_obj.paragraphs:

            if regex.search(p.text):
                inline = p.runs
                # Loop added to work with runs (strings with same style)

                for i in range(len(inline)):
                    if '{' in inline[i - 1].text:
                        changing_field = inline[i].text
                        if changing_field not in list_of_record_headings:
                            list_of_record_headings.append(changing_field)
                            print(changing_field)
                            print("inside if")
                        else:
                            pass
                        print(list_of_record_headings)
        for columns in range(1,len(list_of_record_headings)+1):
            self.sheet1.cell(1,columns).value = list_of_record_headings[columns-1]
            print(self.sheet1.cell(1,1).value)
            print("inside for")
            print(self.sheet1)
        return self.sheet1


class runing_files:


    def create_resource_file(self):
        filename = '../miniproject/Template.docx'
        doc = Document(filename)
        print("inside create resource")
        print("inside playword")
        wb = openpyxl.load_workbook('../miniproject/Book1.xlsx')             #first run this and then
        sh1 = wb['Record1']                                                  #comment these
        print("before creating heading")
        sh1 = Creating_record_file(doc,sh1).Creating_record_headings()       #lines
        wb.save('../miniproject/Book1.xlsx')                                                #this also


    def create_personalized_attachments(self):
        filename = '../miniproject/Template.docx'
        doc = Document(filename)
        print("inside personalize")
        wbnew = openpyxl.load_workbook('../miniproject/Book1.xlsx')
        sh2 = wbnew['Record1']                                             #remove this comment and after commenting above
        replacement(sh2,doc).creating_dictionary()                         #lines
        print("after replacement")
        replacement.updating_attachment_list(attachment_list)









